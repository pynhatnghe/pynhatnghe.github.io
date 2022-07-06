from docx import Document

document = Document()

# Thêm nội dung
document.add_heading('Document Title', 0)
document.add_heading('Chương 1', level=1)
# Có thể dùng add_paragraph sau đó chỉ định style
document.add_paragraph(
    "Kiến thức Python cơ bản",
    style="Heading 2"
)

document.add_paragraph("Loop", style="List Bullet")
document.add_paragraph("List", style="List Bullet")
document.add_paragraph("Dictionary", style="List Bullet")

table = document.add_table(rows=1, cols=3)
table.style = "Table Grid"
cells = table.rows[0].cells
cells[0].text = "STT"
cells[1].text = "Họ tên"
cells[2].text = "Số điện thoại"

# Save xuống
document.save("DemoWord.docx")

# Convert to DOCX to PDF
from docx2pdf import convert
import os

pdf_file = os.path.join(os.getcwd(), "DemoPDF.pdf")
print(pdf_file)
convert("DemoWord.docx", pdf_file)