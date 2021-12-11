from docx2pdf import convert
from docx import Document
from docx.shared import Inches
import os

# Tạo mới file document
document = Document()

document.add_heading("BẢNG ĐIỂM", 0)  # level 0: Title
document.add_heading("Họ tên: Lê Li La", level=2)  # level 2: Heading 2
document.add_paragraph("Phone: 0909123123", style="Heading 3")
document.add_paragraph("Python", style="List Bullet")
document.add_paragraph("SQLite", style="List Bullet")
document.add_paragraph("Tkinter", style="List Bullet")

document.add_paragraph("Liverpool", style="List Number")
document.add_paragraph("Man City", style="List Number")
document.add_paragraph("Chease", style="List Number")

marks = [
    {
        "mon": "Chính trị", "diem": 5, "sotc": 2
    },
    {
        "mon": "Tiếng Anh", "diem": 8, "sotc": 5
    },
    {
        "mon": "Kỹ năng sống", "diem": 9, "sotc": 8
    }
]

table = document.add_table(rows=1, cols=3)
table.style = 'Table Grid'
# Lấy danh sách các cell của dòng 1 (chỉ số đánh 0)
cells = table.rows[0].cells
cells[0].text = "Tên môn"
cells[1].text = "Số tín chỉ"
cells[2].text = "Điểm"

for mark in marks:
    new_row_cells = table.add_row().cells
    new_row_cells[0].text = mark["mon"]
    new_row_cells[1].text = str(mark["diem"])
    new_row_cells[2].text = str(mark["sotc"])

document.save('bangdiem.docx')


# Để convert docx to word cần cài: pip install docx2pdf
# from docx2pdf import convert
current_dir = os.getcwd()
convert(os.path.join(current_dir, 'bangdiem.docx'),
        os.path.join(current_dir, 'bangdiem.pdf'))
